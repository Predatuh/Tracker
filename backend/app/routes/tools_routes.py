"""Tools routes: file converter, PDF merge, image compress."""
import io
import os
import csv
import zipfile
from flask import Blueprint, request, jsonify, send_file
from werkzeug.utils import secure_filename

bp = Blueprint('tools', __name__, url_prefix='/api/tools')

MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB

# ---------------------------------------------------------------------------
# Conversion map: source_ext -> {target_ext: handler_function}
# ---------------------------------------------------------------------------

IMAGE_EXTENSIONS = {'jpg', 'jpeg', 'png', 'bmp', 'tiff', 'tif', 'webp', 'gif'}
SUPPORTED_CONVERSIONS = {}  # populated below after handlers are defined


def _norm_ext(filename):
    """Return lowercase extension without dot, e.g. 'jpg'."""
    return (os.path.splitext(filename)[1] or '').lstrip('.').lower()


def _check_size(file_storage):
    file_storage.seek(0, 2)
    size = file_storage.tell()
    file_storage.seek(0)
    if size > MAX_FILE_SIZE:
        return False
    return True


# ---------------------------------------------------------------------------
# Image <-> Image
# ---------------------------------------------------------------------------

def _convert_image_to_image(file_storage, target_ext):
    from PIL import Image
    img = Image.open(file_storage)
    if img.mode == 'RGBA' and target_ext in ('jpg', 'jpeg', 'bmp'):
        bg = Image.new('RGB', img.size, (255, 255, 255))
        bg.paste(img, mask=img.split()[3])
        img = bg
    elif img.mode != 'RGB' and target_ext in ('jpg', 'jpeg', 'bmp'):
        img = img.convert('RGB')
    buf = io.BytesIO()
    save_fmt = {
        'jpg': 'JPEG', 'jpeg': 'JPEG', 'png': 'PNG', 'bmp': 'BMP',
        'tiff': 'TIFF', 'tif': 'TIFF', 'webp': 'WEBP', 'gif': 'GIF',
    }.get(target_ext, target_ext.upper())
    img.save(buf, format=save_fmt)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Image -> PDF
# ---------------------------------------------------------------------------

def _convert_image_to_pdf(file_storage, _target_ext=None):
    from PIL import Image
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas as rl_canvas
    img = Image.open(file_storage)
    if img.mode == 'RGBA':
        bg = Image.new('RGB', img.size, (255, 255, 255))
        bg.paste(img, mask=img.split()[3])
        img = bg
    elif img.mode != 'RGB':
        img = img.convert('RGB')
    img_buf = io.BytesIO()
    img.save(img_buf, format='PNG')
    img_buf.seek(0)
    pdf_buf = io.BytesIO()
    w, h = img.size
    page_w, page_h = letter
    scale = min(page_w / w, page_h / h, 1.0)
    draw_w, draw_h = w * scale, h * scale
    c = rl_canvas.Canvas(pdf_buf, pagesize=(draw_w, draw_h))
    from reportlab.lib.utils import ImageReader
    c.drawImage(ImageReader(img_buf), 0, 0, draw_w, draw_h)
    c.showPage()
    c.save()
    pdf_buf.seek(0)
    return pdf_buf


# ---------------------------------------------------------------------------
# PDF -> Image
# ---------------------------------------------------------------------------

def _convert_pdf_to_image(file_storage, target_ext):
    from pdf2image import convert_from_bytes
    pages = convert_from_bytes(file_storage.read())
    if not pages:
        raise ValueError('Could not extract any pages from this PDF')
    if len(pages) == 1:
        buf = io.BytesIO()
        fmt = 'JPEG' if target_ext in ('jpg', 'jpeg') else 'PNG'
        pages[0].save(buf, format=fmt)
        buf.seek(0)
        return buf
    # Multiple pages → zip of images
    zip_buf = io.BytesIO()
    fmt = 'JPEG' if target_ext in ('jpg', 'jpeg') else 'PNG'
    ext = 'jpg' if fmt == 'JPEG' else 'png'
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for i, page in enumerate(pages):
            img_buf = io.BytesIO()
            page.save(img_buf, format=fmt)
            zf.writestr(f'page_{i + 1}.{ext}', img_buf.getvalue())
    zip_buf.seek(0)
    return zip_buf, 'zip'


# ---------------------------------------------------------------------------
# DOCX -> PDF
# ---------------------------------------------------------------------------

def _convert_docx_to_pdf(file_storage, _target_ext=None):
    from docx import Document
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

    doc = Document(file_storage)
    pdf_buf = io.BytesIO()
    pdf_doc = SimpleDocTemplate(pdf_buf, pagesize=letter,
                                leftMargin=0.75*inch, rightMargin=0.75*inch,
                                topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    heading_style = ParagraphStyle('DocxHeading', parent=styles['Heading1'], fontSize=14, spaceAfter=8)
    body_style = ParagraphStyle('DocxBody', parent=styles['Normal'], fontSize=11, leading=14, spaceAfter=6)

    story = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            story.append(Spacer(1, 6))
            continue
        # Escape HTML entities for reportlab
        text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        if para.style and para.style.name and 'heading' in para.style.name.lower():
            story.append(Paragraph(text, heading_style))
        else:
            story.append(Paragraph(text, body_style))

    # Also include table content
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = ' | '.join(cells)
            if row_text:
                row_text = row_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(row_text, body_style))

    if not story:
        story.append(Paragraph('(Empty document)', body_style))

    pdf_doc.build(story)
    pdf_buf.seek(0)
    return pdf_buf


# ---------------------------------------------------------------------------
# XLSX -> CSV
# ---------------------------------------------------------------------------

def _convert_xlsx_to_csv(file_storage, _target_ext=None):
    from openpyxl import load_workbook
    wb = load_workbook(file_storage, read_only=True, data_only=True)
    ws = wb.active
    buf = io.StringIO()
    writer = csv.writer(buf)
    for row in ws.iter_rows(values_only=True):
        writer.writerow(['' if cell is None else str(cell) for cell in row])
    result = io.BytesIO(buf.getvalue().encode('utf-8-sig'))
    result.seek(0)
    return result


# ---------------------------------------------------------------------------
# CSV -> XLSX
# ---------------------------------------------------------------------------

def _convert_csv_to_xlsx(file_storage, _target_ext=None):
    from openpyxl import Workbook
    raw = file_storage.read()
    text = raw.decode('utf-8-sig', errors='replace')
    reader = csv.reader(io.StringIO(text))
    wb = Workbook()
    ws = wb.active
    for row in reader:
        ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Build conversion map
# ---------------------------------------------------------------------------

for _src in IMAGE_EXTENSIONS:
    SUPPORTED_CONVERSIONS[_src] = {}
    for _tgt in IMAGE_EXTENSIONS:
        if _tgt != _src and not (_src in ('jpg', 'jpeg') and _tgt in ('jpg', 'jpeg')):
            SUPPORTED_CONVERSIONS[_src][_tgt] = _convert_image_to_image
    SUPPORTED_CONVERSIONS[_src]['pdf'] = _convert_image_to_pdf

SUPPORTED_CONVERSIONS['pdf'] = {
    'jpg': _convert_pdf_to_image,
    'png': _convert_pdf_to_image,
}
SUPPORTED_CONVERSIONS['docx'] = {'pdf': _convert_docx_to_pdf}
SUPPORTED_CONVERSIONS['xlsx'] = {'csv': _convert_xlsx_to_csv}
SUPPORTED_CONVERSIONS['csv'] = {'xlsx': _convert_csv_to_xlsx}


MIME_TYPES = {
    'jpg': 'image/jpeg', 'jpeg': 'image/jpeg', 'png': 'image/png',
    'bmp': 'image/bmp', 'gif': 'image/gif', 'tiff': 'image/tiff', 'tif': 'image/tiff',
    'webp': 'image/webp', 'pdf': 'application/pdf',
    'csv': 'text/csv', 'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'zip': 'application/zip',
}


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@bp.route('/supported-formats', methods=['GET'])
def supported_formats():
    """Return {source_ext: [target_ext, ...]} for the UI to show available conversions."""
    result = {}
    for src, targets in SUPPORTED_CONVERSIONS.items():
        result[src] = sorted(targets.keys())
    return jsonify({'success': True, 'data': result}), 200


@bp.route('/convert', methods=['POST'])
def convert_file():
    """Convert an uploaded file to target_format and return it as a download."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': 'No file selected'}), 400
    if not _check_size(file):
        return jsonify({'error': 'File is too large (max 50 MB)'}), 413

    target_format = (request.form.get('target_format') or '').strip().lower()
    if not target_format:
        return jsonify({'error': 'target_format is required'}), 400

    src_ext = _norm_ext(secure_filename(file.filename))
    if src_ext not in SUPPORTED_CONVERSIONS:
        return jsonify({'error': f'Unsupported source format: .{src_ext}'}), 400
    if target_format not in SUPPORTED_CONVERSIONS[src_ext]:
        return jsonify({'error': f'Cannot convert .{src_ext} to .{target_format}'}), 400

    handler = SUPPORTED_CONVERSIONS[src_ext][target_format]
    try:
        result = handler(file, target_format)
    except Exception as exc:
        return jsonify({'error': f'Conversion failed: {str(exc)}'}), 500

    # Handler may return (buf, actual_ext) for multi-page PDF->image (zip)
    actual_ext = target_format
    if isinstance(result, tuple):
        result, actual_ext = result

    base_name = os.path.splitext(secure_filename(file.filename))[0]
    download_name = f'{base_name}.{actual_ext}'
    mime = MIME_TYPES.get(actual_ext, 'application/octet-stream')
    return send_file(result, mimetype=mime, as_attachment=True, download_name=download_name)


@bp.route('/merge-pdfs', methods=['POST'])
def merge_pdfs():
    """Merge multiple uploaded PDFs into one."""
    from PyPDF2 import PdfMerger
    files = request.files.getlist('files')
    if not files or len(files) < 2:
        return jsonify({'error': 'Upload at least 2 PDF files to merge'}), 400

    merger = PdfMerger()
    try:
        for f in files:
            if not f.filename:
                continue
            ext = _norm_ext(secure_filename(f.filename))
            if ext != 'pdf':
                return jsonify({'error': f'{f.filename} is not a PDF'}), 400
            if not _check_size(f):
                return jsonify({'error': f'{f.filename} is too large (max 50 MB)'}), 413
            merger.append(f)
        buf = io.BytesIO()
        merger.write(buf)
        buf.seek(0)
    except Exception as exc:
        return jsonify({'error': f'Merge failed: {str(exc)}'}), 500
    finally:
        merger.close()

    return send_file(buf, mimetype='application/pdf', as_attachment=True, download_name='merged.pdf')


@bp.route('/compress-image', methods=['POST'])
def compress_image():
    """Compress an uploaded image with a given quality."""
    from PIL import Image
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': 'No file selected'}), 400
    if not _check_size(file):
        return jsonify({'error': 'File is too large (max 50 MB)'}), 413

    ext = _norm_ext(secure_filename(file.filename))
    if ext not in IMAGE_EXTENSIONS:
        return jsonify({'error': f'Unsupported image format: .{ext}'}), 400

    quality = min(max(int(request.form.get('quality', 70)), 1), 100)

    try:
        img = Image.open(file)
        if img.mode == 'RGBA':
            bg = Image.new('RGB', img.size, (255, 255, 255))
            bg.paste(img, mask=img.split()[3])
            img = bg
        elif img.mode != 'RGB':
            img = img.convert('RGB')
        buf = io.BytesIO()
        img.save(buf, format='JPEG', quality=quality, optimize=True)
        buf.seek(0)
    except Exception as exc:
        return jsonify({'error': f'Compression failed: {str(exc)}'}), 500

    base_name = os.path.splitext(secure_filename(file.filename))[0]
    return send_file(buf, mimetype='image/jpeg', as_attachment=True, download_name=f'{base_name}_compressed.jpg')
